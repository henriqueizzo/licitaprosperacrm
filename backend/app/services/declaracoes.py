"""Geração de declarações de habilitação em Word (.docx) com a identidade Prospera.

O checklist da análise IA lista as declarações exigidas pelo edital; aqui o CRM
redige o texto da declaração (IA do provedor ativo, com modelo genérico como
fallback) e monta o .docx: logotipo, título, corpo justificado, linha de data e
bloco de assinatura do representante legal (dados da aba Perfil).
"""
import logging
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from ..models import Licitacao, PerfilEmpresa

logger = logging.getLogger(__name__)

AZUL_PROSPERA = RGBColor(0x25, 0x63, 0xEB)
CINZA_TEXTO = RGBColor(0x33, 0x41, 0x55)

_LOGO = Path(__file__).resolve().parent.parent / "assets" / "prospera-logo.png"

MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
         "agosto", "setembro", "outubro", "novembro", "dezembro"]

SYSTEM_REDATOR = (
    "Você é advogado especialista em licitações públicas brasileiras (Lei 14.133/2021). "
    "Redija declarações formais de habilitação em nome de empresas licitantes. "
    "Responda APENAS com o corpo da declaração, em português formal, SEM título, SEM "
    "cabeçalho, SEM local/data e SEM bloco de assinatura (o sistema adiciona). "
    "Não use markdown nem placeholders inventados — use exatamente os dados fornecidos; "
    "se algum dado não for fornecido, escreva a frase sem ele em vez de deixar lacunas."
)


def _instrucao_ia(licitacao: Licitacao, perfil: PerfilEmpresa, documento: str, referencia: str) -> str:
    empresa = perfil.razao_social or "PROSPERA BENEFÍCIOS S.A."
    cnpj = f", CNPJ {perfil.cnpj}" if perfil.cnpj else ""
    endereco = f", com sede em {perfil.endereco}" if perfil.endereco else ""
    certame = " / ".join(filter(None, [licitacao.modalidade, licitacao.id_externo]))
    # Sem esta trava a IA "empresta" o CNPJ do ÓRGÃO (embutido no número de
    # controle do PNCP) como se fosse o da declarante — visto em teste real.
    aviso_cnpj = (
        "" if perfil.cnpj else
        "\nATENÇÃO: o CNPJ da declarante NÃO foi fornecido — não mencione CNPJ da "
        "empresa e NÃO use o CNPJ que aparece no número do certame (é o CNPJ do órgão)."
    )
    return (
        f"Redija o corpo da seguinte declaração exigida em edital:\n"
        f"- Declaração exigida: {documento}\n"
        f"- Referência no edital: {referencia or 'não informada'}\n"
        f"- Certame: {certame or 'não informado'} — {licitacao.orgao} "
        f"({licitacao.municipio}/{licitacao.uf})\n"
        f"- Objeto: {(licitacao.objeto or '')[:600]}\n"
        f"- Declarante: {empresa}{cnpj}{endereco}, por seu representante legal.\n"
        f"{aviso_cnpj}\n"
        f"O corpo deve começar identificando a empresa declarante e o certame, e então "
        f"declarar exatamente o que a exigência pede, sob as penas da lei, no padrão "
        f"usual de declarações de habilitação da Lei 14.133/2021. Entre 1 e 4 parágrafos."
    )


def _texto_fallback(licitacao: Licitacao, perfil: PerfilEmpresa, documento: str, referencia: str) -> str:
    """Modelo genérico usado quando a IA está indisponível (cota/erro)."""
    empresa = perfil.razao_social or "PROSPERA BENEFÍCIOS S.A."
    cnpj = f", inscrita no CNPJ sob o nº {perfil.cnpj}" if perfil.cnpj else ""
    endereco = f", com sede em {perfil.endereco}" if perfil.endereco else ""
    certame = " / ".join(filter(None, [licitacao.modalidade, licitacao.id_externo]))
    ref = f" (ref.: {referencia})" if referencia else ""
    return (
        f"A empresa {empresa}{cnpj}{endereco}, por intermédio de seu representante legal, "
        f"DECLARA, sob as penas da lei, para fins de participação no certame "
        f"{certame or 'em referência'}, promovido por {licitacao.orgao} "
        f"({licitacao.municipio}/{licitacao.uf}), que cumpre integralmente a exigência "
        f"do edital{ref}: {documento}."
    )


def redigir_texto(licitacao: Licitacao, perfil: PerfilEmpresa, documento: str, referencia: str) -> tuple[str, str]:
    """Retorna (texto, origem) — origem 'ia' ou 'modelo' (fallback)."""
    from ..analyzer import ErroCotaIA, criar_analisador, provedor_ativo

    if provedor_ativo():
        try:
            texto = criar_analisador().redigir(
                _instrucao_ia(licitacao, perfil, documento, referencia), system=SYSTEM_REDATOR
            )
            if texto:
                return texto, "ia"
        except ErroCotaIA as exc:
            logger.warning("IA indisponível ao redigir declaração (%s) — usando modelo genérico", exc)
        except Exception as exc:
            logger.warning("Falha da IA ao redigir declaração (%s) — usando modelo genérico", exc)
    return _texto_fallback(licitacao, perfil, documento, referencia), "modelo"


def _titulo_documento(documento: str) -> str:
    """Título do arquivo/documento: o nome da declaração em caixa alta."""
    titulo = re.sub(r"\s+", " ", documento).strip()
    return (titulo[:120] or "DECLARAÇÃO").upper()


def gerar_docx(licitacao: Licitacao, perfil: PerfilEmpresa, documento: str, referencia: str,
               texto: str) -> bytes:
    doc = Document()
    secao = doc.sections[0]
    secao.top_margin = Cm(2)
    secao.bottom_margin = Cm(2)
    secao.left_margin = Cm(2.5)
    secao.right_margin = Cm(2.5)

    estilo = doc.styles["Normal"]
    estilo.font.name = "Calibri"
    estilo.font.size = Pt(11)
    estilo.font.color.rgb = CINZA_TEXTO

    # Cabeçalho: logo + razão social/CNPJ em azul Prospera
    if _LOGO.is_file():
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.add_run().add_picture(str(_LOGO), height=Cm(1.1))
    p_empresa = doc.add_paragraph()
    p_empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_empresa.add_run(perfil.razao_social or "PROSPERA BENEFÍCIOS S.A.")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = AZUL_PROSPERA
    linhas_empresa = " · ".join(filter(None, [
        f"CNPJ {perfil.cnpj}" if perfil.cnpj else "",
        perfil.endereco,
    ]))
    if linhas_empresa:
        p_dados = doc.add_paragraph()
        p_dados.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_dados.add_run(linhas_empresa)
        r.font.size = Pt(9)

    doc.add_paragraph()

    # Referência do certame (alinhada à direita, como em ofícios)
    certame = " / ".join(filter(None, [licitacao.modalidade, licitacao.id_externo]))
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_ref.add_run(
        "\n".join(filter(None, [
            f"Ref.: {certame}" if certame else "",
            f"{licitacao.orgao} — {licitacao.municipio}/{licitacao.uf}",
            f"Item do edital: {referencia}" if referencia else "",
        ]))
    )
    r.font.size = Pt(9)

    # Título
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_titulo.space_before = Pt(18)
    r = p_titulo.add_run(_titulo_documento(documento))
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = AZUL_PROSPERA
    doc.add_paragraph()

    # Corpo (parágrafos justificados)
    for paragrafo in [p.strip() for p in texto.split("\n") if p.strip()]:
        p = doc.add_paragraph(paragrafo)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(10)

    # Local e data
    hoje = datetime.now()
    local = perfil.cidade_sede or "____________________"
    p_data = doc.add_paragraph(f"{local}, {hoje.day} de {MESES[hoje.month - 1]} de {hoje.year}.")
    p_data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_data.paragraph_format.space_before = Pt(24)

    # Bloco de assinatura
    p_linha = doc.add_paragraph("_______________________________________")
    p_linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_linha.paragraph_format.space_before = Pt(48)
    nome = perfil.representante_nome or "Dario"
    cargo = perfil.representante_cargo or "CEO — Prospera Benefícios"
    for i, linha in enumerate([nome, cargo, perfil.razao_social or "",
                               f"CNPJ {perfil.cnpj}" if perfil.cnpj else ""]):
        if not linha:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(linha)
        r.bold = i == 0
        r.font.size = Pt(11 if i == 0 else 9)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def nome_arquivo(licitacao: Licitacao, documento: str) -> str:
    """Nome do .docx: declaracao-<slug do documento>-<id da licitação>.docx"""
    slug = re.sub(r"[^a-z0-9]+", "-", _titulo_documento(documento).lower()).strip("-")[:60]
    return f"declaracao-{slug or 'habilitacao'}-lic{licitacao.id}.docx"
